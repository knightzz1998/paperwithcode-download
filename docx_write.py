#! /usr/bin/env python
# -*-coding:utf-8-*-

from docx import Document
from docx.shared import Inches


def write_doc(datas, filename):
    document = Document()
    for data_dict in datas:
        title = data_dict['title']
        abstract = data_dict['abstract']
        zh_abstract = data_dict['zh_abstract']
        paper_url = data_dict['paper_url']
        pdf_url = data_dict['pdf_url']
        code_url = data_dict['code_url']
        document.add_heading(title, 2)
        document.add_heading('摘要', 2)
        document.add_paragraph(abstract)
        document.add_paragraph(zh_abstract)
        document.add_paragraph("文章主页 : " + paper_url)
        document.add_paragraph("PDF : " + pdf_url)
        document.add_paragraph("github : " + code_url)
        document.add_paragraph("")

    document.save(filename)


def demo():
    document = Document()

    document.add_heading('Document Title', 0)

    p = document.add_paragraph('A plain paragraph having some ')
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True

    document.add_heading('Heading, level 1', level=1)
    document.add_paragraph('Intense quote', style='Intense Quote')

    document.add_paragraph(
        'first item in unordered list', style='List Bullet'
    )
    document.add_paragraph(
        'first item in ordered list', style='List Number'
    )

    # document.add_picture('monty-truth.png', width=Inches(1.25))

    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam, spam, eggs, and spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save('demo.docx')
